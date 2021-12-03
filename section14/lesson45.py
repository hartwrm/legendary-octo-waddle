import docx
##open doc
d = doc.Document('/some/page/path/demo.docx')

##document contains list of paragraph objects
d.paragraphs
## first paragraph object
d.paragraphs[0]
##each paragraph object has text member varialbe that contains the strings/doucment content
d.paragraphs[0].txt

p = d.paragraphs[1]

##each paragraph contains run objects - new run starts whenever there is a style change
p.runs
p.runs[0].text
p.runs[1].text

##create new word doc
newDoc = docx.Document()
##method to add new paragprah
newDoc.add_paragraph('Hello this is a paragraph')
newDoc.add_paragraph("This is also a paragraph")
#grab first paragraph
p1 = newDoc.paragraphs[0]
## add a new run
p1.add_run('This is a new run')
#set run to bold
p1.runs[1].bold = True


##get all text inside word doc as a string

def getText(filename):
    doc = docx.Document(filename)
    #empty to store text
    fullText = []
    ##iterate through paragraphs in document
    for para in doc.paragraphs:
        ##add paragraphs to empty object
        fullText.append(para.text)
        ##join together - full string
        return '\n'.join(fullText)
print(getText('/your/page/path/andFileName.docx')
